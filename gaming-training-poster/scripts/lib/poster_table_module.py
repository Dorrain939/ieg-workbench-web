"""
poster_table_module.py
======================

海报"复杂多列表格模块"全功能渲染器（单文件 / 零模板依赖 / 可直接拖进任意海报项目）。

本文件覆盖端到端的三种输入入口：
    1) brief 字典 / JSON         → render_table_module(brief)
    2) Excel 文件 (.xlsx)        → excel_to_brief(path) → render_table_module(...)
    3) 自然语言 (LLM 适配)        → nl_to_brief(text, llm_callable) → render_table_module(...)

并提供：
    - JSON Schema 校验 + 业务规则修正
    - 文本超长自动收敛
    - 渲染前预验算（列数 / 字数 / 跨列范围）
    - 错误兜底，避免脏数据进入渲染层

公共 API
--------
    render_table_module(brief, *, scale=2, canvas_width=1920, validate=True) -> PIL.Image
    brief_to_png(brief, out_path, **kw) -> Path
    excel_to_brief(xlsx_path, *, theme="purple_tech") -> dict
    word_to_brief(docx_path, *, theme="purple_tech", tab_col=None, ...) -> dict
    nl_to_brief(text, llm_callable, *, theme="purple_tech") -> dict
    validate_brief(brief) -> brief                   # 校验 + 修正后返回
    get_schema() -> dict                             # 拿 JSON Schema

依赖
----
    必须: pillow, jinja2, playwright (+ chromium)
    可选: pandas + openpyxl  (excel_to_brief 才需要)
    可选: python-docx        (word_to_brief 才需要)
    可选: jsonschema         (用于严格校验，没有也可降级到内置校验)

CLI
---
    python poster_table_module.py render <brief.json> <out.png> [--scale 2 --width 1920]
    python poster_table_module.py from-excel <table.xlsx> <out.png> [--theme purple_tech]
    python poster_table_module.py from-word <table.docx> <out.png> [--theme purple_tech]
    python poster_table_module.py validate <brief.json>
    python poster_table_module.py schema <out.schema.json>

集成进任意海报项目（最小改动）
------------------------------
    from poster_table_module import render_table_module

    module_img = render_table_module(brief["project_overview"])
    canvas.paste(module_img, (paste_x, y_cursor), module_img)
    y_cursor += module_img.height + module_gap
"""

from __future__ import annotations

import argparse
import asyncio
import io
import json
import re
import sys
import tempfile
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from jinja2 import Environment
from PIL import Image

# ---------------------------------------------------------------------------
# 1. 主题 & 限制常量
# ---------------------------------------------------------------------------

THEMES: Dict[str, Dict[str, str]] = {
    "purple_tech": {
        "border": "rgba(168,85,247,0.55)",
        "title_underline": "linear-gradient(90deg, #a855f7, #c084fc)",
        "star": "#c084fc",
        "notice_icon": "#f5b942",
        "header_bg": "#ece5f5",
        "header_text": "#2d1b4e",
        "deliv_bg": "#d3e4a3",
        "deliv_text": "#2d3b1c",
        "text": "#ffffff",
    },
    "blue_business": {
        "border": "rgba(56,189,248,0.55)",
        "title_underline": "linear-gradient(90deg, #38bdf8, #818cf8)",
        "star": "#818cf8",
        "notice_icon": "#fbbf24",
        "header_bg": "#e0e7f5",
        "header_text": "#1e3a5f",
        "deliv_bg": "#fde68a",
        "deliv_text": "#5b3a1c",
        "text": "#ffffff",
    },
    "orange_vivid": {
        "border": "rgba(251,146,60,0.55)",
        "title_underline": "linear-gradient(90deg, #fb923c, #f59e0b)",
        "star": "#fb923c",
        "notice_icon": "#fbbf24",
        "header_bg": "#fff0e0",
        "header_text": "#5b2a0d",
        "deliv_bg": "#bbf7d0",
        "deliv_text": "#1c3b22",
        "text": "#ffffff",
    },
}

LIMITS = {
    "title": 24,
    "tab": 12,
    "header": 24,
    "objective": 240,
    "deliverable": 24,
    "review_label": 24,
    "intro_lead": 24,
    "intro_body": 220,
    "notice_text": 80,
    "min_cols": 2,
    "max_cols": 12,
    "max_sections": 4,
}


# ---------------------------------------------------------------------------
# 2. JSON Schema（用于校验 / 喂给 LLM）
# ---------------------------------------------------------------------------

SCHEMA: Dict[str, Any] = {
    "$schema": "http://json-schema.org/draft-07/schema#",
    "title": "PosterTableModule",
    "type": "object",
    "additionalProperties": False,
    "required": ["sections"],
    "properties": {
        "title": {"type": "string", "maxLength": LIMITS["title"]},
        "theme": {"enum": list(THEMES.keys())},
        "theme_overrides": {
            "type": "object",
            "description": "覆盖基础主题的颜色变量。可用键：border / title_underline / star / notice_icon / header_bg / header_text / deliv_bg / deliv_text / text",
            "additionalProperties": {"type": "string"},
        },
        "notice": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "icon": {"type": "string"},
                "text": {"type": "string", "maxLength": LIMITS["notice_text"]},
            },
        },
        "sections": {
            "type": "array",
            "minItems": 1,
            "maxItems": LIMITS["max_sections"],
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["columns"],
                "properties": {
                    "tab": {"type": "string", "maxLength": LIMITS["tab"]},
                    "tab_color": {"type": "string"},
                    "intro": {
                        "type": "object",
                        "additionalProperties": False,
                        "properties": {
                            "lead": {"type": "string", "maxLength": LIMITS["intro_lead"]},
                            "body": {"type": "string", "maxLength": LIMITS["intro_body"]},
                        },
                    },
                    "columns": {
                        "type": "array",
                        "minItems": LIMITS["min_cols"],
                        "maxItems": LIMITS["max_cols"],
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["header", "objective"],
                            "properties": {
                                "header": {"type": "string", "maxLength": LIMITS["header"]},
                                "objective": {
                                    "type": "string",
                                    "maxLength": LIMITS["objective"],
                                },
                                "deliverable": {
                                    "type": ["string", "null"],
                                    "maxLength": LIMITS["deliverable"],
                                },
                            },
                        },
                    },
                    "review_label": {"type": "string", "maxLength": LIMITS["review_label"]},
                    "comprehensive": {
                        "anyOf": [
                            {"type": "null"},
                            {
                                "type": "object",
                                "additionalProperties": False,
                                "required": ["label", "span_from", "span_to"],
                                "properties": {
                                    "label": {
                                        "type": "string",
                                        "maxLength": LIMITS["header"],
                                    },
                                    "span_from": {"type": "integer", "minimum": 1},
                                    "span_to": {"type": "integer", "minimum": 1},
                                },
                            },
                        ]
                    },
                },
            },
        },
    },
}


def get_schema() -> Dict[str, Any]:
    """获取 JSON Schema 副本（用于 LLM prompt 或文档生成）。"""
    return json.loads(json.dumps(SCHEMA))


# ---------------------------------------------------------------------------
# 3. 校验 + 业务规则修正
# ---------------------------------------------------------------------------

def _truncate(s: Optional[str], limit: int) -> Optional[str]:
    if s is None:
        return None
    s = str(s).strip()
    if len(s) <= limit:
        return s
    return s[: max(1, limit - 1)] + "…"


def _apply_business_rules(brief: Dict[str, Any]) -> Dict[str, Any]:
    """对结构化 brief 做业务级修正：截断超长、修正越界、补默认值。"""
    brief.setdefault("theme", "purple_tech")
    if brief["theme"] not in THEMES:
        brief["theme"] = "purple_tech"

    if "title" in brief:
        brief["title"] = _truncate(brief["title"], LIMITS["title"])

    if brief.get("notice"):
        brief["notice"]["text"] = _truncate(
            brief["notice"].get("text", ""), LIMITS["notice_text"]
        )
        brief["notice"].setdefault("icon", "★")

    sections: List[Dict[str, Any]] = brief.get("sections", [])
    if len(sections) > LIMITS["max_sections"]:
        sections = sections[: LIMITS["max_sections"]]
        brief["sections"] = sections

    default_tab_colors = ["#d8c6e8", "#bfd9ec", "#f5d6c4", "#c5e0c1"]

    for idx, sec in enumerate(sections):
        sec.setdefault("tab", f"区块{idx + 1}")
        sec["tab"] = _truncate(sec["tab"], LIMITS["tab"])
        sec.setdefault("tab_color", default_tab_colors[idx % len(default_tab_colors)])
        sec.setdefault("review_label", "作业验收")
        sec["review_label"] = _truncate(sec["review_label"], LIMITS["review_label"])

        if sec.get("intro"):
            sec["intro"]["lead"] = _truncate(sec["intro"].get("lead", ""), LIMITS["intro_lead"])
            sec["intro"]["body"] = _truncate(sec["intro"].get("body", ""), LIMITS["intro_body"])

        cols: List[Dict[str, Any]] = sec.get("columns", [])
        if len(cols) > LIMITS["max_cols"]:
            cols = cols[: LIMITS["max_cols"]]
        if len(cols) < LIMITS["min_cols"]:
            raise ValueError(
                f"sections[{idx}].columns 至少需要 {LIMITS['min_cols']} 列，当前 {len(cols)}"
            )
        for col in cols:
            col["header"] = _truncate(col.get("header", ""), LIMITS["header"]) or "未命名"
            col["objective"] = _truncate(col.get("objective", ""), LIMITS["objective"]) or ""
            if col.get("deliverable"):
                col["deliverable"] = _truncate(col["deliverable"], LIMITS["deliverable"])
            else:
                col["deliverable"] = None
        sec["columns"] = cols

        # 修正 comprehensive 越界
        comp = sec.get("comprehensive")
        if comp:
            ncols = len(cols)
            sf = max(1, min(int(comp.get("span_from", 1)), ncols))
            st = max(sf, min(int(comp.get("span_to", ncols)), ncols))
            comp["span_from"] = sf
            comp["span_to"] = st
            comp["label"] = _truncate(comp.get("label", "综合大作业"), LIMITS["header"])
            sec["comprehensive"] = comp

    return brief


def validate_brief(brief: Dict[str, Any]) -> Dict[str, Any]:
    """校验 brief 并修正常见问题。返回修正后的 brief。

    优先使用 jsonschema 严格校验；若未安装，退化到内置最小校验。
    """
    if not isinstance(brief, dict):
        raise TypeError(f"brief 必须是 dict，得到 {type(brief).__name__}")

    # 第一步：业务规则先行（截断、补默认）
    brief = _apply_business_rules(json.loads(json.dumps(brief, ensure_ascii=False)))

    # 第二步：jsonschema 严格校验（可选）
    try:
        import jsonschema  # type: ignore[import-not-found]

        jsonschema.validate(brief, SCHEMA)
    except ImportError:
        # 内置最小校验
        if not brief.get("sections"):
            raise ValueError("brief.sections 不能为空")
    except Exception as e:  # noqa: BLE001
        raise ValueError(f"brief 校验失败: {e}") from e

    return brief


# ---------------------------------------------------------------------------
# 4. HTML 模板（内嵌字符串）
# ---------------------------------------------------------------------------

_HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  html, body {
    background: transparent !important;
    font-family: "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei",
                 "Source Han Sans CN", "Noto Sans CJK SC", sans-serif;
    color: {{ T.text }};
    -webkit-font-smoothing: antialiased;
  }
  body { width: {{ canvas_width }}px; padding: 0; }

  .card {
    position: relative;
    background: rgba(255,255,255,0.02);
    border: 1.5px solid {{ T.border }};
    padding: 56px 56px 56px 56px;
    clip-path: polygon(
      24px 0, calc(100% - 24px) 0, 100% 24px,
      100% calc(100% - 24px), calc(100% - 24px) 100%, 24px 100%,
      0 calc(100% - 24px), 0 24px
    );
  }

  .title-wrap { text-align: center; margin-bottom: 28px; position: relative; }
  .title {
    font-size: 56px; font-weight: 900; letter-spacing: 4px;
    color: {{ T.text }}; display: inline-block; position: relative;
    padding: 0 12px;
  }
  .title::after {
    content: ""; display: block; width: 120px; height: 6px;
    background: {{ T.title_underline }}; border-radius: 3px;
    margin: 14px auto 0;
  }
  .title .star { position: absolute; color: {{ T.star }}; font-size: 22px; top: 8px; }
  .title .star.l { left: -36px; }
  .title .star.r { right: -36px; }

  .notice {
    display: flex; align-items: center; gap: 16px;
    margin: 24px 0 28px 0; padding: 4px 0;
  }
  .notice-icon { font-size: 36px; color: {{ T.notice_icon }}; width: 48px; text-align: center; }
  .notice-text { font-size: 22px; font-weight: 600; color: {{ T.text }}; }

  .section {
    display: grid; grid-template-columns: 88px 1fr;
    gap: 20px; margin-bottom: 36px;
  }
  .section:last-child { margin-bottom: 0; }

  .tab {
    border-radius: 10px; display: flex;
    align-items: center; justify-content: center; text-align: center;
    font-size: 26px; font-weight: 800; color: #2d1b4e;
    line-height: 1.3; white-space: pre-line; letter-spacing: 2px;
    padding: 18px 0; min-height: 200px;
  }

  .section-body { display: flex; flex-direction: column; gap: 18px; }

  .intro { font-size: 19px; line-height: 1.7; color: {{ T.text }}; padding: 6px 4px; }
  .intro .lead { color: {{ T.text }}; font-weight: 700; }

  .table { display: grid; gap: 16px 14px; align-items: start; }
  {% for n in range(2, 13) %}
  .table.cols-{{ n }} { grid-template-columns: 110px repeat({{ n }}, 1fr); }
  {% endfor %}

  .row-label {
    font-size: 18px; font-weight: 700; color: {{ T.text }};
    padding-top: 14px; line-height: 1.4; white-space: pre-line;
  }
  .row-label.review { padding-top: 18px; align-self: start; }

  .header-cell {
    background: {{ T.header_bg }}; color: {{ T.header_text }};
    border-radius: 6px; padding: 14px 10px; text-align: center;
    font-size: 19px; font-weight: 800; line-height: 1.35;
    white-space: pre-line; min-height: 64px;
    display: flex; align-items: center; justify-content: center;
  }

  .obj-cell {
    font-size: 16px; line-height: 1.7; color: {{ T.text }};
    padding: 6px 4px 0 4px; text-align: justify;
  }

  .deliv-cell {
    background: {{ T.deliv_bg }}; color: {{ T.deliv_text }};
    border-radius: 4px; padding: 9px 6px; text-align: center;
    font-size: 15px; font-weight: 700; align-self: end;
  }
  .deliv-empty { align-self: end; height: 1px; }

  .comprehensive {
    background: {{ T.deliv_bg }}; color: {{ T.deliv_text }};
    border-radius: 4px; padding: 12px; text-align: center;
    font-size: 17px; font-weight: 800; margin-top: 4px;
  }
</style>
</head>
<body>
  <div class="card">
    {% if data.title %}
    <div class="title-wrap">
      <h1 class="title">
        <span class="star l">&#x2726;</span>{{ data.title }}<span class="star r">&#x2726;</span>
      </h1>
    </div>
    {% endif %}

    {% if data.notice %}
    <div class="notice">
      <div class="notice-icon">{{ data.notice.icon }}</div>
      <div class="notice-text">{{ data.notice.text }}</div>
    </div>
    {% endif %}

    {% for sec in data.sections %}
    {% set ncols = sec.columns | length %}
    <div class="section">
      <div class="tab" style="background: {{ sec.tab_color or '#d8c6e8' }};">{{ sec.tab }}</div>
      <div class="section-body">
        {% if sec.intro %}
        <div class="intro">
          <span class="lead">{{ sec.intro.lead }}</span>{{ sec.intro.body }}
        </div>
        {% endif %}

        <div class="table cols-{{ ncols }}">
          <div class="row-label">课程框架</div>
          {% for col in sec.columns %}
          <div class="header-cell">{{ col.header }}</div>
          {% endfor %}

          <div class="row-label">学习目标</div>
          {% for col in sec.columns %}
          <div class="obj-cell">{{ col.objective }}</div>
          {% endfor %}

          <div class="row-label review">{{ sec.review_label or '作业验收' }}</div>
          {% for col in sec.columns %}
            {% if col.deliverable %}
            <div class="deliv-cell">{{ col.deliverable }}</div>
            {% else %}
            <div class="deliv-empty"></div>
            {% endif %}
          {% endfor %}
        </div>

        {% if sec.comprehensive %}
        {% set sf = sec.comprehensive.span_from %}
        {% set st = sec.comprehensive.span_to %}
        {% set span = st - sf + 1 %}
        <div class="comprehensive"
             style="margin-left: calc(110px + 14px + ({{ sf - 1 }} * (100% - 110px - 14px - {{ (ncols - 1) * 14 }}px) / {{ ncols }}) + ({{ sf - 1 }} * 14px) );
                    width: calc( ({{ span }} * (100% - 110px - 14px - {{ (ncols - 1) * 14 }}px) / {{ ncols }}) + ({{ span - 1 }} * 14px) );">
          {{ sec.comprehensive.label }}
        </div>
        {% endif %}
      </div>
    </div>
    {% endfor %}
  </div>
</body>
</html>
"""


# ---------------------------------------------------------------------------
# 5. 渲染核心
# ---------------------------------------------------------------------------

def _render_html(brief: Dict[str, Any], canvas_width: int) -> str:
    # 先取基础主题，再用 theme_overrides 做精细覆盖
    theme = dict(THEMES.get(brief.get("theme", "purple_tech"), THEMES["purple_tech"]))
    overrides = brief.get("theme_overrides") or {}
    theme.update({k: v for k, v in overrides.items() if k in theme})
    env = Environment(autoescape=False)
    return env.from_string(_HTML_TEMPLATE).render(
        data=brief, T=theme, canvas_width=canvas_width
    )


async def _html_to_png_bytes(html: str, width: int, scale: int) -> bytes:
    try:
        from playwright.async_api import async_playwright
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError(
            "Playwright 未安装。请运行:\n"
            "  pip install playwright\n"
            "  python -m playwright install chromium"
        ) from e

    with tempfile.NamedTemporaryFile(
        suffix=".html", delete=False, mode="w", encoding="utf-8"
    ) as f:
        f.write(html)
        html_path = Path(f.name)
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            context = await browser.new_context(
                viewport={"width": width, "height": 100},
                device_scale_factor=scale,
            )
            page = await context.new_page()
            await page.goto(f"file://{html_path}")
            await page.evaluate("document.fonts.ready")
            png = await page.screenshot(
                omit_background=True, full_page=True, type="png"
            )
            await browser.close()
        return png
    finally:
        try:
            html_path.unlink()
        except OSError:
            pass


def _run_async(coro):
    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(coro)
    import nest_asyncio  # type: ignore[import-not-found]
    nest_asyncio.apply()
    return asyncio.get_event_loop().run_until_complete(coro)


# ---------------------------------------------------------------------------
# 6. 公共渲染 API
# ---------------------------------------------------------------------------

def render_table_module(
    brief: Dict[str, Any],
    *,
    scale: int = 2,
    canvas_width: int = 1920,
    validate: bool = True,
) -> Image.Image:
    """渲染表格模块为 RGBA 透明底 PIL.Image。

    参数
    ----
    brief        : 模块数据（结构见 SCHEMA）
    scale        : DPI 倍率，2 = 高清
    canvas_width : 模块逻辑宽度（CSS px）
    validate     : 是否做 schema 校验 + 业务规则修正（推荐 True）
    """
    if validate:
        brief = validate_brief(brief)
    elif not brief.get("sections"):
        raise ValueError("brief.sections 不能为空")

    html = _render_html(brief, canvas_width=canvas_width)
    png_bytes = _run_async(_html_to_png_bytes(html, width=canvas_width, scale=scale))
    return Image.open(io.BytesIO(png_bytes)).convert("RGBA")


def brief_to_png(
    brief: Dict[str, Any],
    out_path: str | Path,
    *,
    scale: int = 2,
    canvas_width: int = 1920,
    validate: bool = True,
) -> Path:
    """一站式：渲染并保存为 PNG。"""
    img = render_table_module(
        brief, scale=scale, canvas_width=canvas_width, validate=validate
    )
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(out, "PNG")
    return out


# ---------------------------------------------------------------------------
# 7. Excel 适配器
# ---------------------------------------------------------------------------

def excel_to_brief(
    xlsx_path: str | Path,
    *,
    theme: str = "purple_tech",
) -> Dict[str, Any]:
    """从 Excel 自动生成 brief 字典。

    Excel 约定（推荐排布，缺省也能容错）
    -----------------------------------
    每个 sheet 对应一个 section，sheet 名 = tab 文本（用 | 表示换行）

    sheet 内容：
        A1 = 区块 lead 短语，    B1 = 区块 body 文本（可选）
        A2 = "课程框架",  B2..N2 = 各列 header
        A3 = "学习目标",  B3..N3 = 各列 objective
        A4 = "作业验收",  B4..N4 = 各列 deliverable（空白即无）

    特殊行（任意位置）：
        某行 A 列 = "综合大作业" → 其余列内容会跨满，自动识别 span_from/span_to

    顶部元数据（任意 sheet 都可写，第一个找到为准）：
        通过 _meta sheet 提供（可选）：
            title    = 顶部大标题
            notice   = 顶部通知文本
            theme    = 主题名（覆盖参数）
    """
    try:
        import pandas as pd  # type: ignore[import-not-found]
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError("excel_to_brief 需要安装: pip install pandas openpyxl") from e

    xlsx = Path(xlsx_path)
    if not xlsx.exists():
        raise FileNotFoundError(xlsx)

    xl = pd.ExcelFile(xlsx)
    sheet_names = xl.sheet_names

    brief: Dict[str, Any] = {"theme": theme, "sections": []}

    # 顶部元数据
    if "_meta" in sheet_names:
        meta_df = pd.read_excel(xlsx, sheet_name="_meta", header=None)
        for _, row in meta_df.iterrows():
            key = str(row.iloc[0]).strip().lower() if not pd.isna(row.iloc[0]) else ""
            val = str(row.iloc[1]).strip() if len(row) > 1 and not pd.isna(row.iloc[1]) else ""
            if not key or not val:
                continue
            if key == "title":
                brief["title"] = val
            elif key == "notice":
                brief["notice"] = {"icon": "★", "text": val}
            elif key == "theme" and val in THEMES:
                brief["theme"] = val
        sheet_names = [s for s in sheet_names if s != "_meta"]

    default_tab_colors = ["#d8c6e8", "#bfd9ec", "#f5d6c4", "#c5e0c1"]

    for idx, sheet in enumerate(sheet_names):
        df = pd.read_excel(xlsx, sheet_name=sheet, header=None)
        if df.empty:
            continue

        # 建立 tag → row 映射（左列）
        tag_rows: Dict[str, List[Any]] = {}
        intro_lead = ""
        intro_body = ""

        for r in range(len(df)):
            label = df.iloc[r, 0]
            if pd.isna(label):
                continue
            label = str(label).strip()
            row_data = [
                str(df.iloc[r, c]).strip() if not pd.isna(df.iloc[r, c]) else ""
                for c in range(1, len(df.columns))
            ]
            # 第一行：intro
            if r == 0 and label and any(row_data):
                intro_lead = label
                intro_body = next((x for x in row_data if x), "")
                continue
            tag_rows[label] = row_data

        # 找 header / objective / deliverable / comprehensive 行
        def _match(keys, candidates):
            for k in keys:
                for c in candidates:
                    if c in k:
                        return tag_rows[k]
            return []

        all_keys = list(tag_rows.keys())
        headers = _match(all_keys, ["课程框架", "框架", "主题", "Topic", "Header"])
        objectives = _match(all_keys, ["学习目标", "目标", "内容", "Objective"])
        deliverables = _match(all_keys, ["作业", "产出", "验收", "Deliverable"])
        comprehensive_row = _match(all_keys, ["综合大作业", "综合作业", "综合"])

        # 找 review_label
        review_label = "作业验收"
        for k in all_keys:
            if "作业" in k or "验收" in k:
                review_label = k
                break

        # 列数 = headers 非空数量
        ncols = sum(1 for h in headers if h)
        if ncols < LIMITS["min_cols"]:
            continue

        columns = []
        for i in range(ncols):
            columns.append({
                "header": headers[i] if i < len(headers) else "",
                "objective": objectives[i] if i < len(objectives) else "",
                "deliverable": (deliverables[i] if i < len(deliverables) else "") or None,
            })

        # 综合大作业范围
        comp = None
        if comprehensive_row:
            non_empty_idx = [i for i, v in enumerate(comprehensive_row[:ncols]) if v]
            if non_empty_idx:
                label = next((v for v in comprehensive_row if v), "综合大作业")
                comp = {
                    "label": label,
                    "span_from": non_empty_idx[0] + 1,
                    "span_to": non_empty_idx[-1] + 1,
                }

        section = {
            "tab": sheet.replace("|", "\n"),
            "tab_color": default_tab_colors[idx % len(default_tab_colors)],
            "columns": columns,
            "review_label": review_label,
        }
        if intro_lead or intro_body:
            section["intro"] = {"lead": intro_lead, "body": intro_body}
        if comp:
            section["comprehensive"] = comp

        brief["sections"].append(section)

    if not brief["sections"]:
        raise ValueError(f"未能从 Excel 解析出任何有效 section：{xlsx}")

    return validate_brief(brief)


# ---------------------------------------------------------------------------
# 8. Word 适配器（无需 LLM，纯规则解析）
# ---------------------------------------------------------------------------

# 列角色关键词表（匹配表头文字 → 判定该列的语义角色）
_ROLE_KEYWORDS: Dict[str, List[str]] = {
    "tab":         ["专业", "方向", "类别", "分类", "模块", "阶段", "周", "week", "phase", "category"],
    "header":      ["课程", "主题", "框架", "标题", "名称", "topic", "title", "course", "name", "header"],
    "objective":   ["目标", "内容", "描述", "介绍", "说明", "知识", "技能", "objective", "content", "detail", "desc"],
    "deliverable": ["作业", "产出", "交付", "验收", "成果", "deliverable", "output", "homework", "result"],
}


def _detect_role(text: str) -> Optional[str]:
    """根据关键词匹配判断列角色，返回 tab/header/objective/deliverable 或 None。"""
    t = text.strip().lower()
    for role, keywords in _ROLE_KEYWORDS.items():
        if any(kw in t for kw in keywords):
            return role
    return None


def _cell_text(cell) -> str:
    """提取 Word 单元格的纯文本，合并所有段落，去除首尾空白。"""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _is_merge_start(cell) -> bool:
    """判断单元格是否是跨行合并的起始单元格（用于识别 tab 列的合并区）。"""
    try:
        tc = cell._tc
        vMerge = tc.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge"
        )
        if vMerge is None:
            return True  # 没有 vMerge 属性 = 普通单元格（不参与跨行合并）
        val = vMerge.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
        )
        return val == "restart"  # restart = 跨行合并的起始格
    except Exception:  # noqa: BLE001
        return True


def word_to_brief(
    docx_path: str | Path,
    *,
    theme: str = "purple_tech",
    tab_col: Optional[int] = None,
    header_col: Optional[int] = None,
    objective_col: Optional[int] = None,
    deliverable_col: Optional[int] = None,
    section_per_table: bool = False,
) -> Dict[str, Any]:
    """从 Word (.docx) 文件解析表格，生成 table_module brief。

    参数
    ----
    docx_path         : Word 文件路径
    theme             : 基础配色主题（purple_tech / blue_business / orange_vivid）
    tab_col           : 手动指定「区块名」列索引（0-based）；None = 自动推断
    header_col        : 手动指定「课程标题」列索引；None = 自动推断
    objective_col     : 手动指定「学习目标」列索引；None = 自动推断
    deliverable_col   : 手动指定「交付物」列索引；None = 自动推断
    section_per_table : True = 每张 Word 表格生成一个 section；
                        False = 所有表格合并为同一份 brief 的多个 section

    自动推断逻辑
    -----------
    1. 读取第一行（表头行），用关键词表匹配每列角色
    2. 如果第一列跨行合并（rowspan>1），自动识别为 tab 列
    3. 无法推断时：第一列=tab，第二列=header，第三列=objective，第四列=deliverable
    """
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as e:
        raise RuntimeError("word_to_brief 需要安装: pip install python-docx") from e

    doc = Document(Path(docx_path))
    if not doc.tables:
        raise ValueError(f"Word 文件中未找到任何表格：{docx_path}")

    default_tab_colors = ["#d8c6e8", "#bfd9ec", "#f5d6c4", "#c5e0c1"]
    brief: Dict[str, Any] = {"theme": theme, "sections": []}

    for tbl_idx, table in enumerate(doc.tables):
        rows = table.rows
        if not rows:
            continue

        ncols_raw = len(rows[0].cells)

        # ── 步骤1：推断列角色 ──────────────────────────────────────────────
        role_map: Dict[str, int] = {}  # role -> col_idx

        # 优先使用用户手动指定
        if tab_col is not None:
            role_map["tab"] = tab_col
        if header_col is not None:
            role_map["header"] = header_col
        if objective_col is not None:
            role_map["objective"] = objective_col
        if deliverable_col is not None:
            role_map["deliverable"] = deliverable_col

        # 对未指定的角色，用表头关键词推断
        header_row = rows[0]
        for col_i, cell in enumerate(header_row.cells):
            text = _cell_text(cell)
            role = _detect_role(text)
            if role and role not in role_map:
                role_map[role] = col_i

        # 兜底默认映射（保证至少有 header + objective）
        assigned = set(role_map.values())
        defaults = [i for i in range(ncols_raw) if i not in assigned]
        for fallback_role in ["tab", "header", "objective", "deliverable"]:
            if fallback_role not in role_map and defaults:
                role_map[fallback_role] = defaults.pop(0)

        # ── 步骤2：收集剩余「额外列」，拼入 objective ─────────────────────
        known_cols = set(role_map.values())
        extra_cols = [i for i in range(ncols_raw) if i not in known_cols]

        # ── 步骤3：识别 tab 列的跨行分组 ─────────────────────────────────
        # 跳过表头行，从第1行开始
        data_rows = list(rows[1:]) if len(rows) > 1 else []

        # 把 data_rows 按 tab 列的合并区分组
        groups: List[List[Any]] = []
        current_tab = ""
        current_rows: List[Any] = []

        tab_ci = role_map.get("tab")

        for row in data_rows:
            cells = row.cells
            if not cells:
                continue

            if tab_ci is not None and tab_ci < len(cells):
                cell_val = _cell_text(cells[tab_ci])
                is_start = _is_merge_start(cells[tab_ci])
                if is_start and cell_val:
                    # 新的 tab 区块开始
                    if current_rows:
                        groups.append((current_tab, current_rows))
                    current_tab = cell_val
                    current_rows = [cells]
                else:
                    current_rows.append(cells)
            else:
                # 没有 tab 列，把整张表当一个 group
                current_rows.append(cells)

        if current_rows:
            groups.append((current_tab, current_rows))

        # 如果没有识别出任何分组（表格无 tab 列），整表一个 section
        if not groups:
            groups = [("", [row.cells for row in data_rows])]

        # ── 步骤4：把每个分组转成 section ────────────────────────────────
        for sec_idx, (tab_text, group_rows) in enumerate(groups):
            columns = []

            for row_cells in group_rows:
                h_ci = role_map.get("header")
                o_ci = role_map.get("objective")
                d_ci = role_map.get("deliverable")

                header_val = _cell_text(row_cells[h_ci]) if h_ci is not None and h_ci < len(row_cells) else ""
                obj_val    = _cell_text(row_cells[o_ci]) if o_ci is not None and o_ci < len(row_cells) else ""
                deliv_val  = _cell_text(row_cells[d_ci]) if d_ci is not None and d_ci < len(row_cells) else ""

                # 额外列内容拼入 objective
                for ec in extra_cols:
                    if ec < len(row_cells):
                        extra_text = _cell_text(row_cells[ec])
                        if extra_text:
                            obj_val = f"{obj_val}\n{extra_text}".strip() if obj_val else extra_text

                if not header_val and not obj_val:
                    continue  # 跳过全空行

                columns.append({
                    "header":      header_val or "—",
                    "objective":   obj_val,
                    "deliverable": deliv_val or None,
                })

            if len(columns) < LIMITS["min_cols"]:
                continue  # 列数不足，跳过

            section: Dict[str, Any] = {
                "tab":       tab_text or f"区块{sec_idx + 1}",
                "tab_color": default_tab_colors[sec_idx % len(default_tab_colors)],
                "columns":   columns,
                "review_label": "作业验收",
            }
            brief["sections"].append(section)

    if not brief["sections"]:
        raise ValueError(f"未能从 Word 文件解析出有效表格数据：{docx_path}")

    return validate_brief(brief)


# ---------------------------------------------------------------------------
# 8b. 自然语言适配器
# ---------------------------------------------------------------------------

NL_SYSTEM_PROMPT = """你是海报"复杂表格模块"的结构化助手。任务：把用户自然语言需求转成符合给定 JSON Schema 的 JSON。

铁律：
1) 严格遵守 JSON Schema 字段名和字符上限，超长内容必须由你压缩。
2) 不允许新增 schema 之外的字段。
3) 用户提供的关键词（课程名、模块名、人名）必须逐字保留，不许同义改写。
4) 列数推断要忠实于用户描述（说"6 个课程"→ columns 长度必须为 6）。
5) 缺失信息用 null 或空字符串，不要编造。
6) 主题推断：训练营/培训/学习 → purple_tech；认证/项目/汇报 → blue_business；活动/工作坊 → orange_vivid。
7) 只输出 JSON 对象本身，不要任何解释、Markdown 包裹或 ``` 代码块。

JSON Schema:
{schema_json}

用户需求:
{user_text}
"""


def nl_to_brief(
    text: str,
    llm_callable: Callable[[str], str],
    *,
    theme: Optional[str] = None,
    max_repair: int = 1,
) -> Dict[str, Any]:
    """通过自然语言生成 brief。

    参数
    ----
    text         : 用户自然语言描述
    llm_callable : 你的 LLM 调用函数。签名: (full_prompt: str) -> str (返回纯 JSON 字符串)
                   例如:
                       def my_llm(prompt):
                           return openai.chat.completions.create(
                               model="gpt-4o", messages=[{"role":"user","content":prompt}],
                               response_format={"type":"json_object"}
                           ).choices[0].message.content
    theme        : 强制覆盖主题（None 时让 LLM 推断）
    max_repair   : 校验失败时的自动修复重试次数
    """
    schema_str = json.dumps(get_schema(), ensure_ascii=False, indent=2)
    prompt = NL_SYSTEM_PROMPT.format(schema_json=schema_str, user_text=text)

    last_err: Optional[Exception] = None
    raw = ""
    for attempt in range(max_repair + 1):
        full_prompt = prompt
        if attempt > 0 and last_err is not None:
            full_prompt += f"\n\n上次输出错误：{last_err}\n请修正后重新输出 JSON。"

        raw = llm_callable(full_prompt)
        # 抽取 JSON（容错 ```json ``` 包裹）
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if not m:
            last_err = ValueError("LLM 输出未找到 JSON 对象")
            continue
        try:
            brief = json.loads(m.group(0))
            if theme:
                brief["theme"] = theme
            return validate_brief(brief)
        except Exception as e:  # noqa: BLE001
            last_err = e
            continue

    raise RuntimeError(f"nl_to_brief 失败：{last_err}\n最后一次原始输出：{raw[:500]}")


# ---------------------------------------------------------------------------
# 9. CLI
# ---------------------------------------------------------------------------

def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="poster_table_module",
        description="海报复杂表格模块渲染器",
    )
    sub = p.add_subparsers(dest="cmd", required=True)

    pr = sub.add_parser("render", help="从 brief.json 渲染 PNG")
    pr.add_argument("brief")
    pr.add_argument("out")
    pr.add_argument("--scale", type=int, default=2)
    pr.add_argument("--width", type=int, default=1920)
    pr.add_argument("--no-validate", action="store_true")

    pe = sub.add_parser("from-excel", help="从 Excel 渲染 PNG")
    pe.add_argument("xlsx")
    pe.add_argument("out")
    pe.add_argument("--scale", type=int, default=2)
    pe.add_argument("--width", type=int, default=1920)
    pe.add_argument("--theme", default="purple_tech", choices=list(THEMES.keys()))
    pe.add_argument("--save-brief", help="同时保存中间 brief.json")

    pw = sub.add_parser("from-word", help="从 Word (.docx) 渲染 PNG")
    pw.add_argument("docx")
    pw.add_argument("out")
    pw.add_argument("--scale", type=int, default=2)
    pw.add_argument("--width", type=int, default=1920)
    pw.add_argument("--theme", default="purple_tech", choices=list(THEMES.keys()))
    pw.add_argument("--save-brief", help="同时保存中间 brief.json")
    pw.add_argument("--tab-col",         type=int, default=None, help="区块名列索引（0-based）")
    pw.add_argument("--header-col",      type=int, default=None, help="课程标题列索引")
    pw.add_argument("--objective-col",   type=int, default=None, help="学习目标列索引")
    pw.add_argument("--deliverable-col", type=int, default=None, help="交付物列索引")

    pv = sub.add_parser("validate", help="校验 brief.json 是否合法")
    pv.add_argument("brief")

    ps = sub.add_parser("schema", help="导出 JSON Schema")
    ps.add_argument("out")

    return p


def main(argv: Optional[List[str]] = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)

    if args.cmd == "render":
        brief = json.loads(Path(args.brief).read_text(encoding="utf-8"))
        out = brief_to_png(
            brief, args.out,
            scale=args.scale, canvas_width=args.width,
            validate=not args.no_validate,
        )
        img = Image.open(out)
        print(f"[OK] {out}  size={img.size}  mode={img.mode}")
        return 0

    if args.cmd == "from-excel":
        brief = excel_to_brief(args.xlsx, theme=args.theme)
        if args.save_brief:
            Path(args.save_brief).write_text(
                json.dumps(brief, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            print(f"[OK] brief saved -> {args.save_brief}")
        out = brief_to_png(brief, args.out, scale=args.scale, canvas_width=args.width)
        img = Image.open(out)
        print(f"[OK] {out}  size={img.size}  mode={img.mode}")
        return 0

    if args.cmd == "from-word":
        brief = word_to_brief(
            args.docx,
            theme=args.theme,
            tab_col=args.tab_col,
            header_col=args.header_col,
            objective_col=args.objective_col,
            deliverable_col=args.deliverable_col,
        )
        if args.save_brief:
            Path(args.save_brief).write_text(
                json.dumps(brief, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            print(f"[OK] brief saved -> {args.save_brief}")
        out = brief_to_png(brief, args.out, scale=args.scale, canvas_width=args.width)
        img = Image.open(out)
        print(f"[OK] {out}  size={img.size}  mode={img.mode}")
        return 0

    if args.cmd == "validate":
        brief = json.loads(Path(args.brief).read_text(encoding="utf-8"))
        fixed = validate_brief(brief)
        print("[OK] 校验通过。修正后 brief（前 400 字）：")
        print(json.dumps(fixed, ensure_ascii=False, indent=2)[:400])
        return 0

    if args.cmd == "schema":
        Path(args.out).write_text(
            json.dumps(get_schema(), ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(f"[OK] schema -> {args.out}")
        return 0

    parser.print_help()
    return 1


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))



# ---------------------------------------------------------------------------
# 10. Word 精确合并单元格还原渲染（word_table_to_png）  v3
# ---------------------------------------------------------------------------

# ── 共享解析函数（供 word_to_brief 和 word_table_to_png 共用）──────────────

def _parse_word_grid_xml(table) -> tuple:
    """用 XML 直接解析表格，返回 (grid, nrows, ncols)。
    grid: dict {(r,c): {'text', 'rowspan', 'colspan', 'is_cont'}}
    占位续格的 is_cont=True，渲染时跳过。
    """
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def tag(n): return f"{{{NS}}}{n}"

    tbl_xml = table._tbl
    tr_list = tbl_xml.findall(tag("tr"))
    nrows = len(tr_list)

    # 第一遍：确定最大逻辑列数（考虑 gridSpan）
    max_cols = 0
    for tr in tr_list:
        s = 0
        for tc in tr.findall(tag("tc")):
            tcPr = tc.find(tag("tcPr"))
            cs = 1
            if tcPr is not None:
                gs = tcPr.find(tag("gridSpan"))
                if gs is not None:
                    cs = int(gs.get(tag("val"), 1))
            s += cs
        max_cols = max(max_cols, s)
    ncols = max_cols

    grid: Dict[tuple, Dict] = {}
    # occupied: (r,c) -> (origin_r, origin_c)，被 rowspan 占据的格子
    occupied: Dict[tuple, tuple] = {}
    # rowspan_origin: c -> (origin_r, origin_c)，当前列正在延续的 rowspan
    rowspan_origin: Dict[int, tuple] = {}

    for r_idx, tr in enumerate(tr_list):
        tc_list = tr.findall(tag("tc"))
        c_cursor = 0

        for tc in tc_list:
            # 跳过被上方 rowspan 占据的列（不清除 rowspan_origin，只跳游标）
            while c_cursor < ncols and (r_idx, c_cursor) in occupied:
                c_cursor += 1

            if c_cursor >= ncols:
                break

            tcPr = tc.find(tag("tcPr"))
            colspan = 1
            vmerge = "normal"
            if tcPr is not None:
                gs = tcPr.find(tag("gridSpan"))
                if gs is not None:
                    colspan = int(gs.get(tag("val"), 1))
                vm = tcPr.find(tag("vMerge"))
                if vm is not None:
                    vm_val = vm.get(tag("val"), "")
                    vmerge = "restart" if vm_val == "restart" else "cont"

            text = "\n".join(
                "".join(run.text or "" for run in p.findall(".//" + tag("t")))
                for p in tc.findall(".//" + tag("p"))
            ).strip()

            if vmerge == "cont":
                # 续格：找到 origin，rowspan+1，并标记本行所有跨列为 occupied
                if c_cursor in rowspan_origin:
                    or_, oc = rowspan_origin[c_cursor]
                    grid[(or_, oc)]["rowspan"] += 1
                for cc in range(c_cursor, c_cursor + colspan):
                    occupied[(r_idx, cc)] = rowspan_origin.get(cc, (r_idx, c_cursor))
                    grid[(r_idx, cc)] = {"text": "", "rowspan": 1, "colspan": 1, "is_cont": True}
                c_cursor += colspan
                continue

            # 正常格或 restart 格
            cell_entry = {"text": text, "rowspan": 1, "colspan": colspan, "is_cont": False}
            grid[(r_idx, c_cursor)] = cell_entry

            if vmerge == "restart":
                # 登记为 rowspan 起点
                for cc in range(c_cursor, c_cursor + colspan):
                    rowspan_origin[cc] = (r_idx, c_cursor)
            else:
                # 普通格：清除该列的 rowspan 追踪（如果有残留）
                for cc in range(c_cursor, c_cursor + colspan):
                    rowspan_origin.pop(cc, None)

            # colspan 内部占位
            for cc in range(c_cursor + 1, c_cursor + colspan):
                grid[(r_idx, cc)] = {"text": "", "rowspan": 1, "colspan": 1, "is_cont": True}

            c_cursor += colspan

    return grid, nrows, ncols


class _EmptyElem:
    """XML 节点查找返回 None 时的哑对象，避免 None.find() 报错。"""
    def find(self, *a, **kw): return None
    def get(self, *a, **kw): return a[1] if len(a) > 1 else None


def _detect_row_semantic(row_label: str) -> str:
    """根据第1列文字判断行语义类型。"""
    t = row_label.strip()
    if any(k in t for k in ["课程框架", "框架", "主题", "课程名"]):
        return "header"
    if any(k in t for k in ["学习目标", "目标", "内容", "描述"]):
        return "content"
    if any(k in t for k in ["作业", "验收", "交付", "产出"]):
        return "deliv"
    if any(k in t for k in ["其他活动", "活动", "拓展"]):
        return "other"
    return "content"


def _semantic_wrap(text: str) -> str:
    """语义换行：保留原始段落+序号/项目符号前换行，不按标点强制截断。
    对普通内容格，依赖 CSS word-break:keep-all 自然换行，
    只在序号/项目符号前显式换行，保持段落完整性。
    """
    import re
    if not text:
        return text
    raw_paras = [p.strip() for p in text.split("\n") if p.strip()]
    seq_pattern = re.compile(
        r"(?=(?:[①②③④⑤⑥⑦⑧⑨⑩]|[一二三四五六七八九十]+[、．]"
        r"|\d+[、\.．]\s*|[•\-\*]\s*|（\d+）|\(\d+\)))"
    )
    result = []
    for para in raw_paras:
        parts = [p.strip() for p in seq_pattern.split(para) if p.strip()]
        result.extend(parts if parts else [para])
    return "\n".join(result)


def _fix_punct_leading(html_text: str) -> str:
    """禁止标点符号作为行首字符（行首禁排）。
    1. 对已含 <br> 的 HTML 文本，把 <br> 后紧跟的标点移到 <br> 前一行末尾。
    2. 对 CSS 自然换行导致的行首标点，用 U+202F（窄不换行空格）把标点
       和前一个字绑住，阻止在该标点前换行。
    处理的标点：，。、；：！？
    """
    import re
    puncts = r'[，。、；：！？]'
    # 处理显式 <br> 后的标点行首
    html_text = re.sub(r'(<br\s*/?>)(' + puncts + r')', r'\2\1', html_text)
    # 把"字+标点"绑住防止 CSS 自然换行在标点前断行：
    # 在"非空白字符 + 中文标点"之间插入 U+2060（词连接符，零宽不换行）
    WJ = "\u2060"  # WORD JOINER，零宽不换行
    html_text = re.sub(
        r'([^\s<>])(' + puncts + r')',
        lambda m: m.group(1) + WJ + m.group(2),
        html_text,
    )
    return html_text


# 设计自动化策略（固化）：
#
#   列宽自适应
#     - 列0（区块标签）: 固定窄列，竖排文字
#     - 列1（行语义标签）: 固定中宽列
#     - 列2+: 按各列最大内容字数与表头字数*1.5取较大值，用 sqrt(字数) 权重归一化分配
#             最小列宽 160px，确保短文本列不被挤压
#             差值补到最宽列（避免压缩最小列）
#
#   字号自适应
#     - 内容格字符数 > 50:  font-size 缩小到 body_font-2 (text-md)
#     - 内容格字符数 > 100: font-size 缩小到 body_font-4 (text-sm)
#
#   对齐规则（固化铁律）
#     - 「学习目标」类内容行: vertical-align middle + text-align justify + text-align-last left
#                             （多行两端对齐，最后行左对齐，防止中文字间拉伸）
#     - 「课程框架」表头行:   text-align center + vertical-align middle（水平垂直双居中）
#     - 「作业验收」行:       text-align center + vertical-align middle（水平垂直双居中）
#     - 「综合大作业」行:     text-align center + vertical-align middle（水平垂直双居中）
#     - 行语义标签（列1）:    text-align center + vertical-align middle（水平垂直双居中）
#     - 其他活动行:          text-align center + vertical-align middle（badge居中）
#
#   换行规则（固化铁律）
#     - word-break: keep-all —— 中文词不被强制拆断
#     - _semantic_wrap: 序号/项目符号前换行；长段落（>25字）在，；后辅助换行
#     - _nowrap_protect: 中文括号（...）内容不拆行
#     - _fix_punct_leading: 禁止，。、；：！？等标点出现在行首
#     - 表头长文本在 & 符号后插入 <br> 换行
#
#   斑马纹
#     - 学习目标行的奇偶列轻微交替背景（±5% 亮度）
#
#   表头渐变
#     - 课程框架行表头格: linear-gradient(135deg, header_bg → header_bg_dark)
#     - 列0标签格: 竖向渐变 label_bg_top → label_bg_bottom
#
#   「其他活动」行
#     - 统一用轻量 tag 样式（圆角 badge），水平排列
#     - 最后空格子隐藏边框，填充 other_bg 颜色
#
#   整体卡片
#     - 大圆角容器 + 内嵌微发光阴影
#     - 表格最外框较粗（2px），内部格线更细（1px 低透明度）
# ---------------------------------------------------------------------------

import math as _math

_WORD_TABLE_HTML_V3 = r"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  html, body {
    background: transparent !important;
    font-family: "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei",
                 "Source Han Sans CN", "Noto Sans CJK SC", sans-serif;
    -webkit-font-smoothing: antialiased;
  }
  body { width: {{ canvas_width }}px; padding: {{ pad }}px; }

  /* ── 外层卡片 ── */
  .wt-card {
    border-radius: 16px;
    overflow: hidden;
    border: 1px solid {{ T.border }};
    box-shadow: 0 2px 16px {{ T.glow }};
  }

  table {
    width: 100%;
    border-collapse: collapse;
    table-layout: fixed;
  }

  td {
    border: 1px solid {{ T.cell_border }};
    vertical-align: middle;
    word-break: keep-all;
    overflow-wrap: break-word;
    line-break: strict;
    line-height: 1.75;
    white-space: pre-line;
    padding: 10px 12px;
  }

  /* ── 列0：区块大标签（横排，纯色背景，水平+垂直双居中） ── */
  .c-label {
    background: {{ T.label_bg }};
    color: {{ T.label_text }};
    font-size: {{ label_font }}px;
    font-weight: 900;
    letter-spacing: 2px;
    text-align: center;
    vertical-align: middle;
    width: {{ label_col_w }}px;
    min-width: {{ label_col_w }}px;
    border-right: 2px solid {{ T.border }};
    line-height: 1.6;
  }

  /* ── 列1：行语义标签（居中） ── */
  .c-rowlabel {
    background: {{ T.row_label_bg }};
    color: {{ T.row_label_text }};
    font-size: {{ row_label_font }}px;
    font-weight: 700;
    width: {{ row_label_col_w }}px;
    min-width: {{ row_label_col_w }}px;
    border-right: 1px solid {{ T.border }};
    padding: 10px 10px;
    text-align: center;
    vertical-align: middle;
  }

  /* ── 课程框架表头格（扁平纯色，水平+垂直居中） ── */
  /* 铁律：表头水平+垂直双居中，内有圆角白块浮卡效果 */
  .c-header {
    background: {{ T.row_label_bg }};
    color: {{ T.header_text }};
    font-size: {{ header_font }}px;
    font-weight: 700;
    text-align: center;
    vertical-align: middle;
    letter-spacing: 0.3px;
    padding: 8px 6px;
  }
  .c-header .hd-inner {
    background: {{ T.header_bg }};
    border-radius: 8px;
    padding: 10px 10px;
    display: flex;
    align-items: center;
    justify-content: center;
    min-height: 72px;
    width: calc(100% - 4px);
    text-align: center;
    box-sizing: border-box;
  }

  /* ── 学习目标内容格（垂直居中，左对齐，字号自适应） ── */
  /* 铁律：内容格左对齐+垂直居中；字号按字数自动缩放 */
  .c-content {
    background: {{ T.content_bg }};
    color: {{ T.text }};
    font-size: {{ body_font }}px;
    text-align: left;
    vertical-align: middle;
    padding: 12px 14px;
    line-height: 1.8;
  }
  .c-content.text-md { font-size: {{ body_font_md }}px; }
  .c-content.text-sm { font-size: {{ body_font_sm }}px; }

  /* 斑马纹（奇数内容列略深） */
  .c-content.zebra {
    background: {{ T.content_bg_zebra }};
  }

  /* ── 阶段性作业格（水平+垂直居中，圆角按钮风格内容） ── */
  /* 铁律：作业验收行所有内容格居中对齐 */
  .c-deliv {
    background: {{ T.deliv_bg }};
    color: {{ T.deliv_text }};
    font-size: {{ deliv_font }}px;
    font-weight: 700;
    text-align: center;
    vertical-align: middle;
  }

  /* ── 空格子随行语义着色（不露白底） ── */
  .c-empty-deliv   { background: {{ T.deliv_bg_empty }}; }
  .c-empty-other   { background: {{ T.other_bg }}; }
  .c-empty-content { background: {{ T.content_bg }}; }

  /* ── 综合大作业（跨列，强调色，水平+垂直居中） ── */
  /* 铁律：综合大作业跨列行水平+垂直双居中 */
  .c-comprehensive {
    background: {{ T.comp_bg }};
    color: {{ T.comp_text }};
    font-size: {{ deliv_font }}px;
    font-weight: 700;
    letter-spacing: 2px;
    text-align: center;
    vertical-align: middle;
  }

  /* ── 其他活动行：合并格内 badge flex 均匀分布 ── */
  .c-other-rowlabel {
    background: {{ T.other_label_bg }};
    color: {{ T.other_label_text }};
    font-size: {{ row_label_font }}px;
    font-weight: 700;
    border-right: 1px solid {{ T.border }};
    padding: 10px 10px;
    text-align: center;
    vertical-align: middle;
  }
  .c-other {
    background: {{ T.other_bg }};
    vertical-align: middle;
    text-align: center;
    padding: 14px 16px;
  }
  /* 合并格内多个 badge 用 flex 均匀散布 */
  .c-other-merged {
    background: {{ T.other_bg }};
    vertical-align: middle;
    padding: 14px 8px;
  }
  .c-other-merged .badge-row {
    display: flex;
    flex-direction: row;
    justify-content: space-evenly;
    align-items: center;
    flex-wrap: nowrap;
    gap: 8px;
  }
  .c-other .badge,
  .c-other-merged .badge {
    display: inline-block;
    background: {{ T.other_badge_bg }};
    color: {{ T.other_badge_text }};
    border-radius: 16px;
    padding: 6px 20px;
    font-size: {{ other_font }}px;
    font-weight: 600;
    white-space: nowrap;
  }
  .c-other.empty-other {
    background: {{ T.other_bg }};
  }
</style>
</head>
<body>
<div class="wt-card">
<table>
{{ table_html }}
</table>
</div>
</body>
</html>
"""

_WORD_TABLE_THEMES_V3: Dict[str, Dict[str, str]] = {
    # ── light_clean：参考设计图的浅色扁平风（默认主题）──────────────────────
    "light_clean": {
        # 外框：极淡灰，圆角卡片
        "border":           "#d0cce4",
        "inner_border":     "#e8e4f2",
        "glow":             "rgba(100,80,180,0.08)",
        "cell_border":      "#e2ddf0",
        # 列0标签：浅紫薰衣草背景，深灰文字，横排
        "label_bg":         "#ddd5f4",
        "label_text":       "#2c2640",
        # 列1行标签：同浅紫
        "row_label_bg":     "#ddd5f4",
        "row_label_text":   "#2c2640",
        # 表头课程名格：浅奶油白，深灰文字
        "header_bg":        "#f5f0e8",
        "header_bg_dark":   "#f5f0e8",  # 保留兼容，实际不用渐变
        "header_text":      "#2c2640",
        # 内容格：浅青灰，深灰文字
        "content_bg":       "#edf2f7",
        "content_bg_zebra": "#e4ecf4",
        "text":             "#1e1c2e",
        # 验收：浅绿色
        "deliv_bg":         "#dff0c8",
        "deliv_text":       "#1a3808",
        "deliv_bg_empty":   "#eef7e0",
        # 综合大作业
        "comp_bg":          "#dff0c8",
        "comp_bg_l":        "#dff0c8",  # 兼容
        "comp_bg_r":        "#dff0c8",
        "comp_text":        "#1a3808",
        # 其他活动
        "other_bg":         "#f0ecfa",
        "other_border":     "#d0cce4",
        "other_label_bg":   "#ddd5f4",
        "other_label_text": "#2c2640",
        "other_badge_bg":   "#c8bef0",
        "other_badge_text": "#2c2640",
    },
    # ── purple_tech：原深色科技风主题 ────────────────────────────────────────
    "purple_tech": {
        # 外框
        "border":           "rgba(168,85,247,0.75)",
        "inner_border":     "rgba(168,85,247,0.15)",
        "glow":             "rgba(168,85,247,0.18)",
        "cell_border":      "rgba(168,85,247,0.18)",
        # 列0标签
        "label_bg":         "#4a2080",
        "label_bg_top":     "#4a2080",   # 兼容
        "label_bg_bot":     "#2a1050",
        "label_text":       "#f0e8ff",
        # 列1行标签
        "row_label_bg":     "#f0eaff",
        "row_label_text":   "#3b1f6e",
        # 表头
        "header_bg":        "#ddd0f8",
        "header_bg_dark":   "#c4aff0",
        "header_text":      "#2d1b4e",
        # 内容
        "content_bg":       "#180a38",
        "content_bg_zebra": "#201248",
        "text":             "#e2d9ff",
        # 验收
        "deliv_bg":         "#c5e89a",
        "deliv_text":       "#1a3808",
        "deliv_bg_empty":   "#dff0be",
        # 综合大作业
        "comp_bg":          "#8fce50",
        "comp_bg_l":        "#8fce50",
        "comp_bg_r":        "#a8d977",
        "comp_text":        "#1a3808",
        # 其他活动
        "other_bg":         "#120830",
        "other_border":     "rgba(168,85,247,0.12)",
        "other_label_bg":   "#2e1660",
        "other_label_text": "#c9b8f0",
        "other_badge_bg":   "rgba(168,85,247,0.25)",
        "other_badge_text": "#d8c8ff",
    },
    "blue_business": {
        "border":           "rgba(56,189,248,0.75)",
        "inner_border":     "rgba(56,189,248,0.15)",
        "glow":             "rgba(56,189,248,0.15)",
        "cell_border":      "rgba(56,189,248,0.18)",
        "label_bg":         "#1e4a7f",
        "label_bg_top":     "#1e4a7f",
        "label_bg_bot":     "#0d2a50",
        "label_text":       "#e0f2fe",
        "row_label_bg":     "#e8f4ff",
        "row_label_text":   "#1e3a5f",
        "header_bg":        "#bfdbfe",
        "header_bg_dark":   "#93c5fd",
        "header_text":      "#1e3a5f",
        "content_bg":       "#061525",
        "content_bg_zebra": "#091e35",
        "text":             "#dbeafe",
        "deliv_bg":         "#fde68a",
        "deliv_text":       "#5b3a1c",
        "deliv_bg_empty":   "#fef3c7",
        "comp_bg":          "#f59e0b",
        "comp_bg_l":        "#f59e0b",
        "comp_bg_r":        "#fbbf24",
        "comp_text":        "#5b3a1c",
        "other_bg":         "#040e1c",
        "other_border":     "rgba(56,189,248,0.12)",
        "other_label_bg":   "#0f2a46",
        "other_label_text": "#7dd3fc",
        "other_badge_bg":   "rgba(56,189,248,0.2)",
        "other_badge_text": "#bae6fd",
    },
    "orange_vivid": {
        "border":           "rgba(251,146,60,0.75)",
        "inner_border":     "rgba(251,146,60,0.15)",
        "glow":             "rgba(251,146,60,0.15)",
        "cell_border":      "rgba(251,146,60,0.18)",
        "label_bg":         "#8c4010",
        "label_bg_top":     "#8c4010",
        "label_bg_bot":     "#5b2a0d",
        "label_text":       "#fff7ed",
        "row_label_bg":     "#fff7ed",
        "row_label_text":   "#7c3a0d",
        "header_bg":        "#fed7aa",
        "header_bg_dark":   "#fdba74",
        "header_text":      "#5b2a0d",
        "content_bg":       "#1c0a00",
        "content_bg_zebra": "#241200",
        "text":             "#fff1d6",
        "deliv_bg":         "#bbf7d0",
        "deliv_text":       "#1c3b22",
        "deliv_bg_empty":   "#d1fae5",
        "comp_bg":          "#6ee7a0",
        "comp_bg_l":        "#6ee7a0",
        "comp_bg_r":        "#86efac",
        "comp_text":        "#1c3b22",
        "other_bg":         "#110400",
        "other_border":     "rgba(251,146,60,0.12)",
        "other_label_bg":   "#3d1500",
        "other_label_text": "#fed7aa",
        "other_badge_bg":   "rgba(251,146,60,0.22)",
        "other_badge_text": "#ffedd5",
    },
    # ── poster_match：与海报游戏风完全一致的配色 ──────────────────────────
    # 颜色从 vfx_bootcamp_2026 海报像素采样，实现视觉统一
    "poster_match": {
        # 外框：亮紫色边框（对标海报表格圆角边框）
        "border":           "rgba(168,85,247,0.8)",
        "inner_border":     "rgba(168,85,247,0.2)",
        "glow":             "rgba(168,85,247,0.25)",
        "cell_border":      "rgba(168,85,247,0.25)",
        # 列0标签：中深紫，白字（对标海报行标签风格）
        "label_bg":         "#2d1b6e",
        "label_bg_top":     "#3a2280",
        "label_bg_bot":     "#200f55",
        "label_text":       "#f0e8ff",
        # 列1行标签：深紫，白字
        "row_label_bg":     "#1e0f4a",
        "row_label_text":   "#e8d8ff",
        # 表头课程名格：亮紫（海报表头 #9b30f5 系），白字
        "header_bg":        "#7c3aed",
        "header_bg_dark":   "#6d28d9",
        "header_text":      "#ffffff",
        # 内容格：极淡紫白（海报内容行 #f5eaf9），深紫字
        "content_bg":       "#f0e8fb",
        "content_bg_zebra": "#e8d8f5",
        "text":             "#1e0f4a",
        # 阶段性作业/验收：浅绿（保持语义区分）
        "deliv_bg":         "#d4f5c4",
        "deliv_text":       "#1a3808",
        "deliv_bg_empty":   "#e8f9da",
        # 综合大作业
        "comp_bg":          "#bef0a0",
        "comp_bg_l":        "#bef0a0",
        "comp_bg_r":        "#d0f5b4",
        "comp_text":        "#1a3808",
        # 其他活动行：深紫底（海报背景色），白字
        "other_bg":         "#1e0f4a",
        "other_border":     "rgba(168,85,247,0.2)",
        "other_label_bg":   "#2d1b6e",
        "other_label_text": "#e8d8ff",
        "other_badge_bg":   "rgba(168,85,247,0.35)",
        "other_badge_text": "#f0e8ff",
    },
    # ── dark_ref：像素级复刻参考图深紫色主题 ─────────────────────────────────
    # 颜色来源：Pillow 像素采样参考图 /Users/dorrain/Desktop/企业微信截图_f3b9986a...png
    "dark_ref": {
        # 外框：略亮于外背景的深紫边框
        "border":           "#5a4878",
        "inner_border":     "#4a3a68",
        "glow":             "rgba(90,72,120,0.3)",
        "cell_border":      "#4a3a68",
        # 列0标签：浅灰紫背景（参考图表1实测 #c4bbd2），深色文字
        # 注：表2为浅青蓝 #b8d4e0，可通过 theme_overrides 覆盖 label_bg
        "label_bg":         "#c4bbd2",
        "label_bg_top":     "#c4bbd2",
        "label_bg_bot":     "#b0a6c2",
        "label_text":       "#1a1428",
        # 列1行标签：深紫（同外背景色），白色文字
        "row_label_bg":     "#3c2e54",
        "row_label_text":   "#ffffff",
        # 表头课程名格：中灰蓝（实测 #b5bbc8），浅蓝白文字
        "header_bg":        "#b5bbc8",
        "header_bg_dark":   "#aab0be",
        "header_text":      "#1a1e2e",
        # 内容格：深紫（实测 #53466a），浅紫白文字
        "content_bg":       "#53466a",
        "content_bg_zebra": "#4e4164",
        "text":             "#d5d0ea",
        # 阶段性作业/验收：橄榄绿（实测 #b5bba6），近黑文字
        "deliv_bg":         "#b5bba6",
        "deliv_text":       "#0b0b0b",
        "deliv_bg_empty":   "#c8ccbb",
        # 综合大作业：同阶段作业色系略深
        "comp_bg":          "#a8b09a",
        "comp_bg_l":        "#a8b09a",
        "comp_bg_r":        "#b5bba6",
        "comp_text":        "#0b0b0b",
        # 其他活动行：深紫（同行标签背景）
        "other_bg":         "#3c2e54",
        "other_border":     "#4a3a68",
        "other_label_bg":   "#4a3a68",
        "other_label_text": "#d5d0ea",
        "other_badge_bg":   "#6a5888",
        "other_badge_text": "#e0d8f0",
    },
}


def _calc_col_widths(
    grid: Dict,
    nrows: int,
    ncols: int,
    canvas_width: int,
    pad: int,
    label_col_w: int,
    row_label_col_w: int,
) -> List[int]:
    """根据各内容列的最大字符数，用 sqrt 权重自适应分配列宽。"""
    available = canvas_width - 2 * pad - label_col_w - row_label_col_w - (ncols - 2)  # 减去格线1px

    # 每列最大内容字数（跳过标签列）
    col_chars: Dict[int, int] = {}
    for (r, c), v in grid.items():
        if c < 2 or v.get("is_cont"):
            continue
        n = len(v.get("text", ""))
        col_chars[c] = max(col_chars.get(c, 0), n)

    content_cols = sorted(col_chars.keys())
    if not content_cols:
        return [label_col_w, row_label_col_w] + [available // max(1, ncols - 2)] * (ncols - 2)

    # 同时考虑表头行文字长度：表头行(row 0)各列字数取较大值，权重加成 1.5
    header_chars: Dict[int, int] = {}
    for (r, c), v in grid.items():
        if r == 0 and c >= 2 and not v.get("is_cont"):
            header_chars[c] = len(v.get("text", ""))

    # sqrt 权重：正文字数与表头字数*1.5 取大，下限 = 1.0（防止某列宽度太窄）
    effective_chars = {
        c: max(col_chars.get(c, 1), int(header_chars.get(c, 0) * 1.5))
        for c in content_cols
    }
    raw_weights = {c: max(1.0, _math.sqrt(effective_chars[c])) for c in content_cols}
    total_w = sum(raw_weights.values())
    min_col_w = 160  # 最小列宽，确保中文短文本不被挤压
    widths = [label_col_w, row_label_col_w]
    for ci in range(2, ncols):
        if ci in raw_weights:
            w = int(available * raw_weights[ci] / total_w)
        else:
            w = int(available / max(1, len(content_cols)))
        widths.append(max(min_col_w, w))

    # 微调：让总和精确等于 available（均摊到最大列，不压缩最小列）
    diff = available - sum(widths[2:])
    if diff != 0 and widths[2:]:
        # 找最宽的列做调整（避免压缩已达最小宽度的列）
        max_idx = widths.index(max(widths[2:]))
        widths[max_idx] += diff

    return widths


def _nowrap_protect(text: str) -> str:
    """对已 html.escape 后的文本，把中文括号内容用 nowrap span 保护。
    防止 word-break:keep-all 在右括号前断行。
    注意：不保护 & 连接词，避免过长 nowrap 块溢出列宽。
    """
    import re as _re
    # 保护中文括号内的短语，如（讲师线下辅导）→ <span nowrap>（讲师线下辅导）</span>
    text = _re.sub(
        r'（([^）]{1,20})）',
        lambda m: f'<span style="white-space:nowrap">（{m.group(1)}）</span>',
        text,
    )
    return text


def _build_word_html_v3(
    grid: Dict,
    nrows: int,
    ncols: int,
    col_widths: List[int],
    body_font: int,
) -> str:
    """组装 HTML 表格，应用全部设计规则。"""
    import html as html_lib
    import re

    body_font_md = body_font - 2
    body_font_sm = body_font - 4

    lines = []

    # ── 预计算 col0 的真实 rowspan ─────────────────────────────────────────
    # docx 里 col0 往往只有第一行有内容（rs=1），但视觉上应跨越 header+content+deliv
    # 行，不含 other 行。这里自动计算：从 col0 有文字的行开始，连续向下计数直到
    # 遇到 "other" 行或表格结束。
    col0_real_rs: Dict[int, int] = {}  # row_idx → 真实 rowspan
    col0_skip: set = set()             # 被跨行覆盖的行（不再渲染 col0 的 td）
    r = 0
    while r < nrows:
        cell0 = grid.get((r, 0))
        if cell0 is None or cell0.get("is_cont"):
            r += 1
            continue
        # 找到一个 col0 有值的行，向下计算能跨多少行（直到 other 行或表格末尾）
        span = 0
        for rr in range(r, nrows):
            sem = _detect_row_semantic(grid.get((rr, 1), {}).get("text", ""))
            if sem == "other":
                break
            span += 1
        if span > 1:
            col0_real_rs[r] = span
            for rr in range(r + 1, r + span):
                col0_skip.add(rr)
        r += span if span > 0 else 1

    # 生成 colgroup 控制列宽
    lines.append("  <colgroup>")
    for w in col_widths:
        lines.append(f'    <col style="width:{w}px">')
    lines.append("  </colgroup>")

    for r_idx in range(nrows):
        row_label_entry = grid.get((r_idx, 1), {})
        row_semantic = _detect_row_semantic(row_label_entry.get("text", ""))

        lines.append("  <tr>")
        for c_idx in range(ncols):
            cell = grid.get((r_idx, c_idx))

            # col0 在 other 行被 is_cont 跳过时，补一个深紫色格填满宽度
            if c_idx == 0 and row_semantic == "other":
                if cell is None or cell.get("is_cont"):
                    lines.append('    <td class="c-other empty-other"></td>')
                    continue

            if cell is None or cell.get("is_cont"):
                continue

            # col0 被跨行覆盖的行跳过渲染
            if c_idx == 0 and r_idx in col0_skip:
                continue

            rs = cell.get("rowspan", 1)
            cs = cell.get("colspan", 1)
            text = cell.get("text", "")

            # col0 用真实计算的 rowspan 覆盖
            if c_idx == 0 and r_idx in col0_real_rs:
                rs = col0_real_rs[r_idx]

            rs_attr = f' rowspan="{rs}"' if rs > 1 else ""
            cs_attr = f' colspan="{cs}"' if cs > 1 else ""

            # ── 决定 CSS class 和样式 ──────────────────────────────
            if c_idx == 0:
                # 列0：区块标签；换行控制：在中文字符与前面内容之间插入 <wbr>
                # 使浏览器优先在"英文/数字段 → 中文"的边界处换行
                # 例："Unity 3D特效" → "Unity 3D<wbr>特效"（在数字后、中文前断）
                css = "c-label"
                import re as _re_label
                escaped = html_lib.escape(text)
                # 在"英文/数字末尾 → 中文开头"处插入 <wbr>
                escaped = _re_label.sub(r'([A-Za-z0-9])([^\x00-\x7F])', r'\1<wbr>\2', escaped)
                # 移除半角空格（避免在空格处额外断行，保持"Unity 3D"作为一组）
                text_html = escaped

            elif c_idx == 1:
                # 列1：行语义标签
                if row_semantic == "other":
                    css = "c-other-rowlabel"
                    text_html = html_lib.escape(text)
                elif row_semantic == "header":
                    # 表头行的行标签格也用 hd-inner 浮卡效果
                    css = "c-rowlabel"
                    text_html = f'<div class="hd-inner">{html_lib.escape(text)}</div>'
                else:
                    css = "c-rowlabel"
                    text_html = html_lib.escape(text)

            elif row_semantic == "header":
                # 课程框架表头行：用 hd-inner div 实现圆角浮卡效果
                css = "c-header"
                wrapped = _semantic_wrap(text)
                # 超长表头：在 & 符号后插入换行，防止 keep-all 不在此处断行
                if len(text) > 12:
                    wrapped = wrapped.replace("&", "&\n")
                escaped = html_lib.escape(wrapped)
                inner_html = _nowrap_protect(escaped).replace("\n", "<br>")
                inner_html = _fix_punct_leading(inner_html)
                text_html = f'<div class="hd-inner">{inner_html}</div>'

            elif row_semantic == "content":
                # 学习目标/普通内容行
                if not text:
                    css = "c-empty-content"
                    text_html = ""
                elif any(k in text for k in ["综合大作业", "综合作业"]) and cs > 1:
                    css = "c-comprehensive"
                    text_html = html_lib.escape(text)
                else:
                    wrapped = _semantic_wrap(text)
                    nchars = len(text)
                    size_class = ""
                    if nchars > 100:
                        size_class = " text-sm"
                    elif nchars > 50:
                        size_class = " text-md"
                    # 斑马纹：内容列索引奇偶
                    zebra = " zebra" if (c_idx % 2 == 1) else ""
                    css = f"c-content{size_class}{zebra}"
                    raw_html = html_lib.escape(wrapped).replace("\n", "<br>")
                    text_html = _fix_punct_leading(raw_html)

            elif row_semantic == "deliv":
                # 作业验收行
                if not text:
                    css = "c-empty-deliv"
                    text_html = ""
                elif any(k in text for k in ["综合大作业", "综合作业"]) and cs > 1:
                    css = "c-comprehensive"
                    text_html = html_lib.escape(text)
                else:
                    css = "c-deliv"
                    escaped = html_lib.escape(text)
                    text_html = _nowrap_protect(escaped).replace("\n", "<br>")

            elif row_semantic == "other":
                # 其他活动行：col2-6 合并为一个大格，内部 flex 均匀分布 badge
                # 策略：遇到 c_idx==2 时，收集本行 col2..ncols-1 所有有文字的格，
                # 输出一个 colspan 合并格 + badge-row；c_idx>2 的格跳过
                if c_idx == 1:
                    # 行标签格正常输出
                    css = "c-other-rowlabel"
                    text_html = html_lib.escape(text)
                elif c_idx == 2:
                    # 收集 col2..ncols-1 所有非空文字
                    badges = []
                    for ci2 in range(2, ncols):
                        cv = grid.get((r_idx, ci2))
                        if cv and not cv.get("is_cont") and cv.get("text", "").strip():
                            badges.append(html_lib.escape(cv["text"].strip()))
                    badge_html = "".join(
                        f'<span class="badge">{b}</span>' for b in badges
                    )
                    merged_cs = ncols - 2  # 合并从 col2 到 ncols-1
                    cs_attr = f' colspan="{merged_cs}"'
                    rs_attr = ""
                    css = "c-other-merged"
                    text_html = f'<div class="badge-row">{badge_html}</div>'
                    lines.append(f'    <td{rs_attr}{cs_attr} class="{css}">{text_html}</td>')
                    continue  # 已手动输出，跳过下面的通用输出
                else:
                    # col>2 的格已被合并，跳过
                    continue

            else:
                css = "c-empty-content"
                text_html = html_lib.escape(text).replace("\n", "<br>") if text else ""

            lines.append(f'    <td{rs_attr}{cs_attr} class="{css}">{text_html}</td>')

        lines.append("  </tr>")

    return "\n".join(lines)


def word_table_to_png(
    docx_path: str | Path,
    out_path: str | Path,
    *,
    table_index: int = 0,
    canvas_width: int = 1200,
    scale: int = 2,
    pad: int = 36,
    label_col_w: int = 90,
    row_label_col_w: int = 130,
    label_font: int = 17,
    row_label_font: int = 13,
    header_font: int = 14,
    body_font: int = 13,
    deliv_font: int = 13,
    other_font: int = 13,
    theme: str = "dark_ref",
    theme_overrides: Optional[Dict[str, str]] = None,
) -> Path:
    """从 Word (.docx) 精确还原表格合并结构，应用自动化设计优化，渲染为高质量 PNG。

    默认主题 dark_ref 为像素级复刻参考图的深紫色风格。
    可传 theme="light_clean" / "purple_tech" 等切换风格。
    海报宽度约束：canvas_width 默认 1200，建议范围 1000~1440（超宽放海报上字会变小）。

    设计自动化内容（无需手动调整，规则已固化）：
    - 列宽：sqrt(字数) 权重分配，最小 160px，表头字数权重 ×1.5
    - 字号：>50字缩一档，>100字缩两档
    - 列0：横排大标签，水平+垂直居中
    - 内容格：顶对齐+左对齐（long text 自然 CSS 换行）
    - 表头/作业验收/综合大作业：水平+垂直双居中
    - 标点行首禁排：，。、；：！？ 不出现在行首
    - 空格子配色随行语义（不露白底）
    - 斑马纹：内容行奇偶列微深交替
    - 其他活动行 badge 圆角 tag 样式
    """
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as e:
        raise RuntimeError("需要安装: pip install python-docx") from e

    doc = Document(Path(docx_path))
    if not doc.tables:
        raise ValueError(f"Word 文件中未找到任何表格：{docx_path}")

    table = doc.tables[table_index]
    grid, nrows, ncols = _parse_word_grid_xml(table)

    col_widths = _calc_col_widths(
        grid, nrows, ncols, canvas_width, pad, label_col_w, row_label_col_w
    )

    table_html = _build_word_html_v3(grid, nrows, ncols, col_widths, body_font)

    T = dict(_WORD_TABLE_THEMES_V3.get(theme, _WORD_TABLE_THEMES_V3["dark_ref"]))
    if theme_overrides:
        T.update({k: v for k, v in theme_overrides.items() if k in T})

    env = Environment(autoescape=False)
    html = env.from_string(_WORD_TABLE_HTML_V3).render(
        T=T,
        canvas_width=canvas_width,
        pad=pad,
        label_col_w=label_col_w,
        row_label_col_w=row_label_col_w,
        label_font=label_font,
        row_label_font=row_label_font,
        header_font=header_font,
        body_font=body_font,
        body_font_md=body_font - 2,
        body_font_sm=body_font - 4,
        deliv_font=deliv_font,
        other_font=other_font,
        table_html=table_html,
    )

    png_bytes = _run_async(_html_to_png_bytes(html, width=canvas_width, scale=scale))
    img = Image.open(io.BytesIO(png_bytes)).convert("RGBA")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    img.save(out, "PNG")
    return out
